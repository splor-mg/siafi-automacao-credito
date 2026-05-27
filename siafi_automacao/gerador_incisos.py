"""
Conferencia de incisos em minutas de decreto.

Le todos os .docx de uma pasta, extrai paragrafos que comecem com
"do saldo" ou "do excesso" e contenham "no valor de", consolida
incisos com a MESMA descricao (somando os valores), gera o extenso
em portugues e salva tudo em conferencia.xlsx.

Requer: python-docx, openpyxl
Instalar com: pip install python-docx openpyxl
"""

import os
import re
import glob
from decimal import Decimal, ROUND_HALF_UP
from docx import Document
from openpyxl import Workbook


# =============================================================================
# CONFIGURACAO - ajuste estes dois caminhos conforme o seu ambiente
# =============================================================================
BASE_DIR = r"C:\Users\x70167581686\OneDrive - CAMG\General\@dcmefo\2026\Decretos\Minutas Finalizadas"


# =============================================================================
# EXTENSO EM PORTUGUES (pt-BR)
# =============================================================================
UNI = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove",
       "dez", "onze", "doze", "treze", "catorze", "quinze", "dezesseis",
       "dezessete", "dezoito", "dezenove"]
DEZ = ["", "", "vinte", "trinta", "quarenta", "cinquenta",
       "sessenta", "setenta", "oitenta", "noventa"]
CEN = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos",
       "seiscentos", "setecentos", "oitocentos", "novecentos"]


def _ate_999(n: int) -> str:
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    c, r = divmod(n, 100)
    partes = []
    if c:
        partes.append(CEN[c])
    if r:
        if r < 20:
            partes.append(UNI[r])
        else:
            d, u = divmod(r, 10)
            partes.append(DEZ[d] if u == 0 else f"{DEZ[d]} e {UNI[u]}")
    return " e ".join(partes)


def _inteiro_extenso(n: int) -> str:
    if n == 0:
        return "zero"
    bilhoes, r = divmod(n, 1_000_000_000)
    milhoes, r = divmod(r, 1_000_000)
    milhares, centenas = divmod(r, 1_000)
    partes = []
    if bilhoes:
        partes.append(f"{_ate_999(bilhoes)} {'bilhão' if bilhoes == 1 else 'bilhões'}")
    if milhoes:
        partes.append(f"{_ate_999(milhoes)} {'milhão' if milhoes == 1 else 'milhões'}")
    if milhares:
        partes.append("mil" if milhares == 1 else f"{_ate_999(milhares)} mil")
    if centenas:
        partes.append(_ate_999(centenas))
    # entre blocos (bilhao/milhao/mil/centenas) usa-se apenas espaco,
    # seguindo o padrao do arquivo de referencia
    return " ".join(partes)


def valor_extenso(valor: float) -> str:
    d = Decimal(str(valor)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    inteiro = int(d)
    centavos = int((d - inteiro) * 100)
    partes = []
    if inteiro:
        partes.append(f"{_inteiro_extenso(inteiro)} {'real' if inteiro == 1 else 'reais'}")
    if centavos:
        nome = "centavo" if centavos == 1 else "centavos"
        ext = f"{_inteiro_extenso(centavos)} {nome}"
        partes.append(f"e {ext}" if partes else ext)
    return " ".join(partes) if partes else "zero reais"


def fmt_valor(valor: float) -> str:
    """Formata 22224054.11 -> '22.224.054,11'."""
    d = Decimal(str(valor)).quantize(Decimal("0.01"), ROUND_HALF_UP)
    s = f"{d:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def romano(n: int) -> str:
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, s in vals:
        while n >= v:
            out += s
            n -= v
    return out


# =============================================================================
# LEITURA DOS .DOCX
# =============================================================================
def ler_paragrafos_docx(caminho: str) -> list[str]:
    """Retorna todos os paragrafos de texto de um .docx (corpo + tabelas)."""
    doc = Document(caminho)
    textos = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    textos.append(p.text)
    return textos


# =============================================================================
# EXTRACAO E CONSOLIDACAO DOS INCISOS
# =============================================================================
PADRAO_FRASE = re.compile(r"(?i)(do\s+saldo|do\s+excesso)")
PADRAO_NO_VALOR = re.compile(r"(?i)no\s+valor")
# captura tudo de "do saldo" / "do excesso" ate (sem incluir) "no valor de"
PADRAO_CHAVE = re.compile(
    r"(?i)(do\s+(?:saldo|excesso).*?)\s*,?\s*no\s+valor\s+de"
)
# captura o valor monetario logo apos "R$"
PADRAO_VALOR = re.compile(r"R\$\s*(\d[\d\.]*,\d{2})")


def parse_valor_br(s: str) -> float:
    """'16.380.103,30' -> 16380103.30"""
    return float(s.replace(".", "").replace(",", "."))


def processar_pasta(pasta: str) -> list[dict]:
    """Le todos os .docx da pasta e retorna lista de dicts {chave, valor}."""
    arquivos = sorted(glob.glob(os.path.join(pasta, "*.docx")))
    arquivos = [a for a in arquivos if not os.path.basename(a).startswith("~$")]

    if not arquivos:
        raise FileNotFoundError(f"Nenhum arquivo .docx encontrado em: {pasta}")

    registros = []
    for caminho in arquivos:
        try:
            paragrafos = ler_paragrafos_docx(caminho)
        except Exception as e:
            print(f"Erro ao ler {caminho}: {e}")
            continue

        for texto in paragrafos:
            # remove espaco entre "R$" e o numero (R$ 16.380 -> R$16.380)
            t = re.sub(r"(R\$)\s+", r"\1", texto)
            # so paragrafos que contenham as palavras-chave E "no valor"
            if not (PADRAO_FRASE.search(t) and PADRAO_NO_VALOR.search(t)):
                continue
            m_chave = PADRAO_CHAVE.search(t)
            m_valor = PADRAO_VALOR.search(t)
            if not m_chave or not m_valor:
                continue
            chave = m_chave.group(1).strip()
            valor = parse_valor_br(m_valor.group(1))
            registros.append({"chave": chave, "valor": valor})

    return registros


def consolidar(registros: list[dict]) -> list[dict]:
    """Agrupa por chave normalizada (case-insensitive + espacos normalizados),
    soma os valores e preserva a primeira capitalizacao encontrada."""
    grupos: dict[str, dict] = {}
    ordem: list[str] = []
    for r in registros:
        chave_norm = re.sub(r"\s+", " ", r["chave"].lower()).strip()
        if chave_norm not in grupos:
            grupos[chave_norm] = {"chave_texto": r["chave"], "valor": 0.0}
            ordem.append(chave_norm)
        grupos[chave_norm]["valor"] += r["valor"]

    saida = []
    for i, chave_norm in enumerate(ordem):
        g = grupos[chave_norm]
        num = romano(i + 2)  # comeca em II como no script original
        valor = round(g["valor"], 2)
        inciso = (
            f"{num} - {g['chave_texto']} no valor de "
            f"R${fmt_valor(valor)} ({valor_extenso(valor)});"
        )
        saida.append({"incisos": inciso, "valor": valor})
    return saida


# =============================================================================
# ESCRITA DO EXCEL
# =============================================================================
def salvar_xlsx(linhas: list[dict], caminho_saida: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "base"
    ws.append(["incisos", "valor"])
    for r in linhas:
        ws.append([r["incisos"], r["valor"]])
    # largura util
    ws.column_dimensions["A"].width = 120
    ws.column_dimensions["B"].width = 18
    wb.save(caminho_saida)


# =============================================================================
# MAIN
# =============================================================================
def main() -> None:
    subpasta = input("Informe a subpasta (ex: 05 maio\\minuta 84): ").strip()
    pasta = os.path.join(BASE_DIR, subpasta)
    registros = processar_pasta(pasta)
    if not registros:
        print("Nenhum inciso de 'do saldo'/'do excesso' com 'no valor de' encontrado.")
        return
    linhas = consolidar(registros)
    saida = os.path.join(pasta, "conferencia.xlsx")
    salvar_xlsx(linhas, saida)
    print(f"Processo concluido. {len(registros)} incisos lidos, "
          f"{len(linhas)} apos consolidacao. Arquivo: {saida}")


if __name__ == "__main__":
    main()